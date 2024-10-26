import streamlit as st
from crewai import Agent, Task, Crew, Process
import os
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
#from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64
#from langchain_community.llms import Ollama
from langchain_google_genai import ChatGoogleGenerativeAI



#llm = Ollama(model = "llama3:8b", max_tokens = 8000)

load_dotenv()

#os.environ["OPENAI_API_KEY"] = "sk-proj-PIwwf9DykQtEWxCXxPn9T3BlbkFJN9bEX0EF5vbOBUZ4nTS2"
#os.environ["SERPER_API_KEY"] = "5aa99b51d7a15965a9beb15c9d6202f02f658123"
os.environ["GOOGLE_API_KEY"] = "AIzaSyD4U4sHWxGNIKe_3sMRr6AaxOf1YFgSp44"
# LLM object and API Key
#os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")
#os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")


# def generate_docx(result):
#     doc = Document()
#     doc.add_heading('Legal Guidance', 0)
#     doc.add_paragraph(result)
#     bio = BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio

# def get_download_link(bio, filename):
#     b64 = base64.b64encode(bio.read()).decode()
#     return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Legal Guidance</a>'

# st.set_page_config(
#     layout="wide"
# )
# # Title
# st.title("ChainVerdict Legal Adviser")

def generate_docx(result):
    # Ensure result is a string
    if result is None:
        result = "No result provided."
    elif not isinstance(result, str):
        try:
            # Convert the object to a string
            result = str(result)
        except Exception as e:
            result = f"Error converting result to string: {str(e)}"

    doc = Document()
    doc.add_heading('Legal Guidance', 0)
    clean_result = result.replace("**", "")
     # Add the cleaned result to the document
    doc.add_paragraph(clean_result)
    # Add the result as a paragraph without stars
    # paragraphs = result.splitlines()  # Split the result into lines
    # for para in paragraphs:
    #     if para.strip():  # Add non-empty lines
    #         doc.add_paragraph(para.strip())  # Add the line as a paragraph without stars or extra spaces

    # Save the document in a BytesIO object
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)  # Reset the stream position to the beginning
    return bio

def get_download_link(bio, filename):
    # Encode the document and create a download link
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Legal Guidance</a>'

st.set_page_config(
    layout="wide"
)

# Title
st.title("ChainVerdict Legal Advisor")


# Input Fields
gender = st.selectbox('Select Gender of the petitioner', ('Male', 'Female', 'Other'))
age = st.number_input('Enter Age of the petitioner', min_value=0, max_value=120, value=25)
case_type = st.selectbox('Select Case Type', ('Criminal', 'Civil', 'Family Law'))
legal_issue = st.text_area('Enter Specific Legal Issue', 'e.g., theft, breach of contract, custody dispute')
allegations_claims = st.text_area('Enter Allegations or Claims', 'e.g., accused of theft, suing for damages')

party_role = st.selectbox('Select Your Role in the Case', ('Petitioner', 'Respondent'))
parties_involved = st.text_area('Enter Parties Involved', 'e.g., John Doe (petitioner), Jane Doe (respondent)')

case_facts = st.text_area('Enter Specific Facts and Circumstances', 'e.g., theft occurred on 12th July 2023 at XYZ location')
incident_details = st.text_area('Enter Incident Details', 'e.g., date, location, witnesses')
evidence_details = st.text_area('Enter Evidence Details', 'e.g., witness statements, CCTV footage')

desired_outcome = st.text_area('Enter Desired Outcome', 'e.g., seeking compensation, custody of child')
relief_sought = st.text_area('Enter Relief Sought', 'e.g., financial compensation, restraining order')


# Initialize Tools
# search_tool = SerperDevTool()
# scrape_tool = ScrapeWebsiteTool()

# llm = ChatOpenAI(
#     model="gpt-3.5-turbo-16k",
#     temperature=0.1,
#     max_tokens=8000
# )
# llm = Ollama(
#     model="llama3:8b"
#     # temperature=0.1,
#     # max_tokens=8000
# )
llm1 = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash",
    temperature=0,
    max_tokens=20000
)
 
llm2 = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro",
    temperature=0,
    max_tokens=20000
)




# Define Agents
case_analyzer = Agent(
    role="Case Analyzer",
    goal="Analyze the case based on the provided details.",
    backstory="This agent specializes in understanding the specifics of a case based on input details. It uses legal knowledge to identify relevant factors and preliminary legal considerations.",
    verbose=True,
    allow_delegation=False,
    #tools=[search_tool, scrape_tool],
    llm=llm1
)

petition_suggester = Agent(
    role="Petition Suggester",
    goal="Suggest possible petitions to be filed against the crime based on the Bharatiya Nyay Sahita (BNS) in India Judiciary and follow https://www.mha.gov.in/sites/default/files/250883_english_01042024.pdf this link and give the user the coreect BNS section according to the crime.",
    backstory="This agent specializes in the legal procedures and documentation required for filing petitions. It provides suggestions based on the latest laws and sections from the judiciary of India. Follow the link https://www.mha.gov.in/sites/default/files/250883_english_01042024.pdf for better result",
    verbose=True,
    allow_delegation=False,
    #tools=[search_tool, scrape_tool],
    llm=llm1
)

next_steps_advisor = Agent(
    role="Next Steps Advisor like a lawyer",
    goal="Advise on the next steps considering the new law with proper BNS (Bharatiya Nyay Sahita) sections and the constitution of India Judiciary Law like a lawyer will suggest.",
    backstory="This agent provides guidance on the legal process following the analysis and petition suggestions. It ensures compliance with current legal standards and procedures and It will provide you all the petitions related to the case like your lawyer.",
    verbose=True,
    allow_delegation=False,
    #tools=[search_tool, scrape_tool],
    llm=llm1
)

# fir_creator = Agent(
#     role="FIR Creator",
#     goal="Generate the analysis  in the FIR format based on the Indian judiciary system.",
#     backstory="This agent creates official documents in the required format for submission. It generates FIRs that can be directly used at police stations.",
#     verbose=True,
#     allow_delegation=False,
#     #tools=[search_tool, scrape_tool],
#     llm=llm1
# )

# Define Tasks
case_analysis_task = Task(
    description=(
        "1. Analyze the case details: age ({age}) of the petitioner, gender ({gender}) of the petitioner, case type ({case_type}), legal issue ({legal_issue}), allegations/claims ({allegations_claims}), "
        "parties involved ({parties_involved}), role ({party_role}), facts ({case_facts}), incident details ({incident_details}), and evidence ({evidence_details}).\n"
        "2. Provide a legal analysis with potential considerations."
    ),
    expected_output="A legal analysis with key considerations.",
    agent=case_analyzer
)

petition_task = Task(
    description=(
        "1. Based on the case analysis, suggest possible petitions to be filed according to the link given.\n"
        "2. Reference the relevant sections of the Bharatiya Nyay Sahita (BNS) of India Judiciary Law."
    ),
    expected_output="A list of possible petitions with BNS sections references.",
    agent=petition_suggester
)

next_steps_task = Task(
    description=(
        "1. Advise on the next steps to be taken based on the case and suggested petitions like a lawyer.\n"
        "2. Include the relevant BNS (Bharatiya Nyay Sahita) sections and laws to ensure compliance with current standards and legal procedures. "
    ),
    expected_output="A step-by-step guide for the next legal steps with BNS sections references.",
    agent=next_steps_advisor
)

# fir_task = Task(
#     description=(
#         "1. write the analysis in the FIR format of Indian Judiciary based on the provided case details.\n"
#         "2. Ensure the FIR adheres to the Indian judiciary system requirements."
#     ),
#     expected_output="A written FIR with the analyzed case in FIR format ready for submission.",
#     agent=fir_creator
# )


# Create Crew
crew = Crew(
    agents=[case_analyzer, petition_suggester, next_steps_advisor], #fir_creator], 
    #agents=[case_analyzer],
    tasks=[case_analysis_task, petition_task, next_steps_task], #fir_task],
    #tasks=[case_analysis_task],
    verbose= True
)



# Execution
if st.button("Generate Legal Guidance"):
    with st.spinner('Processing...'):
        result = crew.kickoff(inputs={
            "age": age, 
            "gender": gender,
            "case_type": case_type,
            "legal_issue": legal_issue,
            "allegations_claims": allegations_claims,
            "parties_involved": parties_involved,
            "party_role": party_role,
            "case_facts": case_facts,
            "incident_details": incident_details,
            "evidence_details": evidence_details,
            "desired_outcome": desired_outcome,
            "relief_sought": relief_sought
        })
        st.write(result)
        docx_file = generate_docx(result)

        download_link = get_download_link(docx_file, "legal_guidance.docx")
        st.markdown(download_link, unsafe_allow_html=True)

