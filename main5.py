import streamlit as st
from crewai import Agent, Task, Crew
import os
from langchain_google_genai import ChatGoogleGenerativeAI
from docx import Document
from io import BytesIO
import base64
from dotenv import load_dotenv

load_dotenv()

# Set API keys
os.environ["GOOGLE_API_KEY"] = "AIzaSyD4U4sHWxGNIKe_3sMRr6AaxOf1YFgSp44"
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

# Define LLMs
llm1 = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash",
    temperature=0,
    max_tokens=20000
)

llm2 = ChatGoogleGenerativeAI(
    model="gemini-1.5-pro",
    temperature=0,
    max_tokens=20000
)

# Function to generate DOCX document
def generate_docx(result):
    if result is None:
        result = "No result provided."
    elif not isinstance(result, str):
        try:
            result = str(result)
        except Exception as e:
            result = f"Error converting result to string: {str(e)}"

    doc = Document()
    doc.add_heading('First Information Report (FIR)', 0)
    clean_result = result.replace("**", "")
    doc.add_paragraph(clean_result)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Function to generate download link
def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Legal Guidance</a>'

st.set_page_config(layout="wide")
st.title("ChainVerdict Legal")

# Initialize session state for interaction flow
if 'step' not in st.session_state:
    st.session_state.step = 0
    st.session_state.case_details = {}

# Step 1: Basic Information
if st.session_state.step == 0:
    st.subheader("Step 1: Basic Case Information")
    date_of_occurrence = st.date_input('Enter Date of Occurrence')
    time_of_occurrence = st.time_input('Enter Time of Occurrence')
    place_of_occurrence = st.text_input('Enter Place of Occurrence (e.g., Street, Area)')
    police_station = st.text_input('Enter Police Station Name')
    district = st.text_input('Enter District Name')
    state = st.text_input('Enter State Name')
    complainant_name = st.text_input('Enter Complainant\'s Name')
    complainant_address = st.text_area('Enter Complainant\'s Address')
    complainant_occupation = st.text_input('Enter Complainant\'s Occupation')

    if st.button("Submit Basic Information"):
        # Store basic information in session state
        st.session_state.case_details.update({
            "date_of_occurrence": date_of_occurrence,
            "time_of_occurrence": time_of_occurrence,
            "place_of_occurrence": place_of_occurrence,
            "police_station": police_station,
            "district": district,
            "state": state,
            "complainant_name": complainant_name,
            "complainant_address": complainant_address,
            "complainant_occupation": complainant_occupation
        })
        st.session_state.step = 1

# Step 2: Request Additional Information
if st.session_state.step == 1:
    st.subheader("Step 2: Provide Additional Case Details")
    st.write("I cannot generate a complete FIR with the provided information. Please provide more details about the case, including:")
    st.markdown("""
    * **Nature of the crime:** Is this theft, misplacement, or something else?
    * **Details of the crime:** Are there any suspects? Was the item stolen, or is it possible it was misplaced?
    * **The victim:** Is there any information about the person who reported the crime?
    * **The accused:** Is there anyone suspected of committing the crime?
    """)

    nature_of_offense = st.text_input('Enter Nature of Offense (e.g., Theft, Assault)')
    details_of_offense = st.text_area('Enter Specific Details of the Offense')
    accused_name = st.text_input('Enter Accused\'s Name (if known)')
    accused_address = st.text_area('Enter Accused\'s Address (if known)')
    evidence = st.text_area('Enter Evidence Details (e.g., Witness Statements, CCTV Footage)')

    if st.button("Submit Additional Information"):
        # Store additional information in session state
        st.session_state.case_details.update({
            "nature_of_offense": nature_of_offense,
            "details_of_offense": details_of_offense,
            "accused_name": accused_name,
            "accused_address": accused_address,
            "evidence": evidence
        })
        st.session_state.step = 2

# Step 3: Generate FIR
if st.session_state.step == 2:
    st.subheader("Step 3: Generate FIR")
    st.write("Generating FIR with the provided details...")
    # Define Agents
    petition_suggester = Agent(
        role="Petition Suggester",
        goal="Suggest possible petitions to be filed against the crime based on the Bharatiya Nyay Sahita (BNS) in India Judiciary and follow https://www.mha.gov.in/sites/default/files/250883_english_01042024.pdf this link and give the user the coreect BNS section according to the crime.",
        backstory="This agent specializes in the legal procedures and documentation required for filing petitions. It provides suggestions based on the latest laws and sections from the judiciary of India. Follow the link https://www.mha.gov.in/sites/default/files/250883_english_01042024.pdf for better result",
        verbose=True,
        allow_delegation=False,
        #tools=[search_tool, scrape_tool],
        llm=llm1
    )
    fir_creator = Agent(
        role="FIR Creator",
        goal="Generate the FIR based on the provided case details and the format should be like the given https://police.py.gov.in/Police%20manual/Forms%20pdf/FORM%20IF%201.pdf link.",
        backstory="This agent creates official FIR documents that include all details provided by the complainant with the suggested petitions by petition suggester and leaves the police actions to be filled by authorities.",
        verbose=True,
        allow_delegation=False,
        llm=llm1
    )
    # Define Tasks
    petition_task = Task(
        description=(
        "1. Based on the case analysis, write all possible petitions to be filed according to the link given.\n"
        "2. Reference the relevant sections of the Bharatiya Nyay Sahita (BNS) of India Judiciary Law."
    ),
    expected_output="A list of possible petitions with BNS sections references.",
    agent=petition_suggester
    )
    fir_task = Task(
    description=(
        "1. Use the provided case details: Date of Occurrence ({date_of_occurrence}), Time of Occurrence ({time_of_occurrence}), "
            "Place of Occurrence ({place_of_occurrence}), Police Station ({police_station}), District ({district}), State ({state}), "
            "Complainant Name ({complainant_name}), Complainant Address ({complainant_address}), Complainant Occupation ({complainant_occupation}), "
            "Accused Name ({accused_name}), Accused Address ({accused_address}), "
            "Nature of Offense ({nature_of_offense}), Details of Offense ({details_of_offense}), Evidence ({evidence}).\n"
            "2. Generate an FIR with these details and follow the given Link https://police.py.gov.in/Police%20manual/Forms%20pdf/FORM%20IF%201.pdf to provide the exact format of FIR to the user."
    ),
    expected_output="A written FIR with the provided details ready for submission.",
    agent=fir_creator
    )
    # Create Crew
    crew = Crew(
    agents=[petition_suggester, fir_creator],
    tasks=[petition_task, fir_task],
    verbose=True
    )
    # Convert date and time to strings
    date_str = st.session_state.case_details["date_of_occurrence"].strftime('%Y-%m-%d')
    time_str = st.session_state.case_details["time_of_occurrence"].strftime('%H:%M:%S')
    
    result = crew.kickoff(inputs={
        "date_of_occurrence": date_str,
        "time_of_occurrence": time_str,
        "place_of_occurrence": st.session_state.case_details["place_of_occurrence"],
        "police_station": st.session_state.case_details["police_station"],
        "district": st.session_state.case_details["district"],
        "state": st.session_state.case_details["state"],
        "complainant_name": st.session_state.case_details["complainant_name"],
        "complainant_address": st.session_state.case_details["complainant_address"],
        "complainant_occupation": st.session_state.case_details["complainant_occupation"],
        "accused_name": st.session_state.case_details["accused_name"],
        "accused_address": st.session_state.case_details["accused_address"],
        "nature_of_offense": st.session_state.case_details["nature_of_offense"],
        "details_of_offense": st.session_state.case_details["details_of_offense"],
        "evidence": st.session_state.case_details["evidence"]
    })

    st.write(result)
    docx_file = generate_docx(result)
    download_link = get_download_link(docx_file, "FIR.docx")
    st.markdown(download_link, unsafe_allow_html=True)

    # Reset the session state after generation
    st.session_state.step = 0
    st.session_state.case_details = {}